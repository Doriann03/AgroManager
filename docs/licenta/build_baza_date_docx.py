from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "23_capitolul_3_baza_de_date_mld_mpd.md"
OUTPUT = ROOT / "Capitol_3_Baza_de_date_tabele.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths_cm):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Cm(widths_cm[idx])


def style_table(table):
    table.style = "Table Grid"
    set_table_width(table, [4.0, 3.0, 2.8, 7.0])
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(10)
            if row_idx == 0:
                set_cell_shading(cell, "D9EAF7")
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.bold = True


def add_md_paragraph(doc, line):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Cm(1)
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(line.replace("`", ""))
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def add_heading(doc, text, level):
    paragraph = doc.add_paragraph()
    paragraph.style = f"Heading {min(level, 3)}"
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.bold = True
    run.font.size = Pt(14 if level == 1 else 12)


def parse_table(lines, start):
    header = [cell.strip() for cell in lines[start].strip("|").split("|")]
    rows = [header]
    idx = start + 2
    while idx < len(lines) and lines[idx].startswith("|"):
        rows.append([cell.strip() for cell in lines[idx].strip("|").split("|")])
        idx += 1
    return rows, idx


def build_docx():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(12)

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    idx = 0
    while idx < len(lines):
        line = lines[idx].strip()
        if not line:
            idx += 1
            continue

        if line.startswith("# "):
            add_heading(doc, line[2:].strip(), 1)
            idx += 1
            continue
        if line.startswith("## "):
            add_heading(doc, line[3:].strip(), 2)
            idx += 1
            continue
        if line.startswith("### "):
            add_heading(doc, line[4:].strip(), 3)
            idx += 1
            continue
        if line.startswith("|") and idx + 1 < len(lines) and lines[idx + 1].startswith("|"):
            rows, idx = parse_table(lines, idx)
            table = doc.add_table(rows=len(rows), cols=len(rows[0]))
            for r, row in enumerate(rows):
                for c, value in enumerate(row):
                    table.cell(r, c).text = value
            style_table(table)
            doc.add_paragraph()
            continue

        add_md_paragraph(doc, line)
        idx += 1

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_docx()
    print(OUTPUT)
